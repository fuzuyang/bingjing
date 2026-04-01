import os
import re
import sys
from typing import List, Optional, Tuple

from sqlalchemy.orm import Session

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.append(PROJECT_ROOT)

from core.database import SessionLocal, init_db
from core.models import SourceDocument, generate_content_hash


class DataIngester:
    """Load files from a folder and persist raw documents to biz_source_documents."""

    CONTENT_COLS = [
        "文章内容",
        "正文",
        "内容",
        "摘要",
        "Content",
        "content",
        "text",
    ]
    TITLE_COLS = [
        "标题",
        "文章标题",
        "name",
        "Name",
        "Title",
        "subject",
    ]
    NOISE_KEYWORDS = ["HTTP ERROR", "502 Bad Gateway", "404 Not Found"]
    SUPPORTED_EXTS = {".csv", ".xlsx", ".xls", ".txt", ".md", ".pdf", ".docx"}

    def __init__(self, data_folder: str, recursive: bool = False):
        self.data_folder = data_folder
        self.recursive = recursive

    def _list_files(self) -> List[str]:
        if not os.path.isdir(self.data_folder):
            return []

        files: List[str] = []
        if self.recursive:
            for root, _, names in os.walk(self.data_folder):
                for name in names:
                    ext = os.path.splitext(name)[1].lower()
                    if ext in self.SUPPORTED_EXTS:
                        files.append(os.path.join(root, name))
        else:
            for name in os.listdir(self.data_folder):
                file_path = os.path.join(self.data_folder, name)
                ext = os.path.splitext(name)[1].lower()
                if os.path.isfile(file_path) and ext in self.SUPPORTED_EXTS:
                    files.append(file_path)
        return sorted(files)

    @staticmethod
    def _clean_text(text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"[\.\-_]{3,}", "", text)
        text = re.sub(r"^\d+\s*$", "", text, flags=re.MULTILINE)
        text = text.replace("&nbsp;", " ").replace("\xa0", " ")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _read_table(self, file_path: str):
        pd = self._load_pandas()
        if pd is None:
            return None

        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".csv":
                try:
                    return pd.read_csv(file_path, encoding="utf-8-sig")
                except UnicodeDecodeError:
                    return pd.read_csv(file_path, encoding="gbk")
            if ext in {".xlsx", ".xls"}:
                return pd.read_excel(file_path)
        except Exception as exc:
            print(f"  ! 读取表格失败: {file_path} -> {exc}")
        return None

    @staticmethod
    def _extract_row(df, row, idx: int, file_stem: str) -> Tuple[str, str]:
        def has_value(v) -> bool:
            return v is not None and str(v).strip().lower() not in {"", "nan", "none"}

        title = next(
            (str(row[c]) for c in DataIngester.TITLE_COLS if c in df.columns and has_value(row[c])),
            "",
        ).strip()
        content = next(
            (str(row[c]) for c in DataIngester.CONTENT_COLS if c in df.columns and has_value(row[c])),
            "",
        ).strip()

        if not title:
            title = f"{file_stem}-{idx + 1}"
        return title, content

    def _extract_from_table(self, file_path: str) -> List[Tuple[str, str]]:
        pd = self._load_pandas()
        if pd is None:
            print(f"  ! 跳过表格(缺少pandas依赖): {file_path}")
            return []

        df = self._read_table(file_path)
        if df is None or df.empty:
            return []

        file_stem = os.path.splitext(os.path.basename(file_path))[0]
        docs: List[Tuple[str, str]] = []
        for idx, row in df.iterrows():
            title, content = self._extract_row(df, row, idx, file_stem)
            docs.append((title, content))
        return docs

    @staticmethod
    def _load_pandas():
        try:
            import pandas as pd
            return pd
        except Exception:
            return None

    @staticmethod
    def _extract_from_text_file(file_path: str) -> Optional[Tuple[str, str]]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                text = f.read()
        except Exception as exc:
            print(f"  ! 读取文本失败: {file_path} -> {exc}")
            return None

        title = os.path.splitext(os.path.basename(file_path))[0]
        return title, text

    @staticmethod
    def _extract_from_pdf(file_path: str) -> Optional[Tuple[str, str]]:
        try:
            from pypdf import PdfReader
        except Exception:
            print(f"  ! 跳过PDF(缺少pypdf依赖): {file_path}")
            return None

        try:
            reader = PdfReader(file_path)
            pages = [page.extract_text() or "" for page in reader.pages]
            content = "\n".join(pages)
            title = os.path.splitext(os.path.basename(file_path))[0]
            return title, content
        except Exception as exc:
            print(f"  ! 读取PDF失败: {file_path} -> {exc}")
            return None

    @staticmethod
    def _extract_from_docx(file_path: str) -> Optional[Tuple[str, str]]:
        try:
            from docx import Document  # Optional dependency
        except Exception:
            print(f"  ! 跳过DOCX(缺少python-docx依赖): {file_path}")
            return None

        try:
            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            content = "\n".join(paragraphs)
            title = os.path.splitext(os.path.basename(file_path))[0]
            return title, content
        except Exception as exc:
            print(f"  ! 读取DOCX失败: {file_path} -> {exc}")
            return None

    def _extract_documents(self, file_path: str) -> List[Tuple[str, str]]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext in {".csv", ".xlsx", ".xls"}:
            return self._extract_from_table(file_path)
        if ext in {".txt", ".md"}:
            one = self._extract_from_text_file(file_path)
            return [one] if one else []
        if ext == ".pdf":
            one = self._extract_from_pdf(file_path)
            return [one] if one else []
        if ext == ".docx":
            one = self._extract_from_docx(file_path)
            return [one] if one else []
        return []

    @staticmethod
    def _infer_category(file_name: str, title: str = "") -> str:
        sample = f"{file_name} {title}".lower()
        speech_keys = ["讲话", "发言", "致辞", "演讲", "speech"]
        case_keys = ["案例", "判决", "裁定", "裁判", "审判", "case", "judgment"]
        if any(key in sample for key in speech_keys):
            return "speech"
        if any(key in sample for key in case_keys):
            return "case"
        return "policy"

    @staticmethod
    def _is_noise(text: str) -> bool:
        if len(text) < 30:
            return True
        return any(k in text for k in DataIngester.NOISE_KEYWORDS)

    @staticmethod
    def _load_existing_hashes(db: Session) -> set:
        rows = db.query(SourceDocument.content_hash).all()
        return {r[0] for r in rows if r[0]}

    def run(self):
        init_db()
        db: Session = SessionLocal()

        try:
            files = self._list_files()
            if not files:
                print(f"未在目录中发现可处理文件: {self.data_folder}")
                return

            print(f"开始入库，目标目录: {self.data_folder}")
            print(f"待处理文件数: {len(files)}")

            processed_hashes = self._load_existing_hashes(db)
            total_new = 0
            total_skipped = 0

            for file_path in files:
                file_name = os.path.basename(file_path)
                print(f"\n- 解析文件: {file_name}")
                docs = self._extract_documents(file_path)
                if not docs:
                    print("  ! 未抽取到有效文档内容")
                    continue

                file_new = 0
                file_skip = 0
                for idx, (raw_title, raw_content) in enumerate(docs):
                    content = self._clean_text(raw_content)
                    if self._is_noise(content):
                        file_skip += 1
                        continue

                    content_hash = generate_content_hash(content)
                    if not content_hash or content_hash in processed_hashes:
                        file_skip += 1
                        continue

                    title = (raw_title or "").strip() or f"{os.path.splitext(file_name)[0]}-{idx + 1}"
                    category = self._infer_category(file_name=file_name, title=title)

                    try:
                        new_doc = SourceDocument(
                            file_name=file_name,
                            title=title[:512],
                            content_text=content,
                            content_hash=content_hash,
                            category=category,
                        )
                        db.add(new_doc)
                        db.commit()
                        processed_hashes.add(content_hash)
                        file_new += 1
                    except Exception as exc:
                        db.rollback()
                        file_skip += 1
                        print(f"  ! 文档写入失败({title[:30]}): {str(exc)[:100]}")

                total_new += file_new
                total_skipped += file_skip
                print(f"  完成: 新增 {file_new} 条, 跳过 {file_skip} 条")

            print("\n入库结束")
            print(f"总新增: {total_new}")
            print(f"总跳过: {total_skipped}")
        finally:
            db.close()


if __name__ == "__main__":
    default_folder = os.path.join(PROJECT_ROOT, "raw_file")
    ingester = DataIngester(default_folder)
    ingester.run()
